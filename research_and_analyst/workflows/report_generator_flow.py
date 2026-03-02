import os
import sys
import re
from datetime import datetime
from typing import Optional
from langgraph.types import Send
from jinja2 import Template
from langgraph.graph import StateGraph, START, END
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_community.tools.tavily_search import TavilySearchResults
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from langgraph.checkpoint.memory import MemorySaver
from research_and_analyst.schemas.models import Perspectives, GenerateAnalystsState, ResearchGraphState
from research_and_analyst.utils.model_loader import ModelLoader
from research_and_analyst.workflows.interview_workflow import InterviewGraphBuilder
from research_and_analyst.prompt_library.prompt_locator import (
    CREATE_ANALYSIS_PROMPT,
    INTRO_CONCLUSION_INSTRUCTIONS,
    REPORT_WRITER_INSTRUCTIONS,
)

from research_and_analyst.log.logger import CustomLogger
from research_and_analyst.exception.custom_exception import ResearchAnalystException


class AutonomousReportGenerator:
    """
    handles end-to-end autonomous report generation
    """

    def __init__(self, llm):
        self.llm = llm
        self.memory = MemorySaver()

        api_key = os.getenv("TAVILY_API_KEY")
        if not api_key:
            raise RuntimeError("TAVILY_API_KEY not set in environment")
        self.tavily_search = TavilySearchResults(tavily_api_key=api_key)

        # ✅ FIX: CustomLogger is a factory; get the structlog logger instance
        self.logger = CustomLogger().get_logger(__file__)

        # Optional: fail fast if logger is wrong type
        if not hasattr(self.logger, "info") or not hasattr(self.logger, "error"):
            raise TypeError(f"Logger misconfigured, got: {type(self.logger)}")

    def create_analyst(self, state: GenerateAnalystsState):
        topic = state["topic"]
        max_analyst = state["max_analysts"]
        human_analyst_feedback = state.get("human_analyst_feedback", "")

        try:
            self.logger.info("initiated persona creating", topic=topic)
            structured_llm = self.llm.with_structured_output(Perspectives)
            system_prompt = CREATE_ANALYSIS_PROMPT.render(
                topic=topic,
                max_analysts=max_analyst,
                human_analyst_feedback=human_analyst_feedback,
            )
            analysts = structured_llm.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content="generate a set of analysts"),
                ]
            )
            return {"analysts": analysts.analysts}

        except Exception as e:
            self.logger.error("Failed to generate set of analysts", error=str(e))
            raise ResearchAnalystException("check your code Bonita!", e)

    def human_feedback(self):
        """
        await node for human feedback
        """
        try:
            self.logger.info("awaiting human feedback")
        except Exception as e:
            self.logger.error("error in feedback stage", error=str(e))
            raise ResearchAnalystException("sorry buddy! you are mistaken", e)

    def write_report(self, state: ResearchGraphState):
        """
        compile all sections into one single comprehensive report
        """
        sections = state.get("sections", [])
        topic = state.get("topic", "")

        try:
            if not sections:
                sections = ["no sections generated - please check"]

            self.logger.info("report writing generated", topic=topic)
            system_prompt = REPORT_WRITER_INSTRUCTIONS.render(topic=topic)
            report = self.llm.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content="\n\n".join(sections)),
                ]
            )
            self.logger.info("Report generated successfully")
            return {"content": report.content}

        except Exception as e:
            self.logger.error("report generation failed", error=str(e))
            raise ResearchAnalystException("check your code darling", e)

    def write_introduction(self, state: ResearchGraphState):
        """
        writes the introduction part of the report
        """
        topic = state["topic"]
        section = state.get("sections", [])
        formatted_str_sections = "\n\n".join([f"{s}" for s in section])

        try:
            self.logger.info("generating the introductions", topic=topic)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic, formatted_str_sections=formatted_str_sections
            )
            intro = self.llm.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content="write the introduction part of the report"),
                ]
            )
            self.logger.info("Introduction generated", length=len(intro.content))
            return {"introduction": intro.content}

        except Exception as e:
            self.logger.error("some error occurred", error=str(e))
            raise ResearchAnalystException("check your code buddy", e)

    def write_conclusion(self, state: ResearchGraphState):
        """
        writes the conclusion of the report
        """
        section = state.get("sections", [])
        topic = state["topic"]

        try:
            self.logger.info("generated the conclusion", topic=topic)
            formatted_str_sections = "\n\n".join([f"{s}" for s in section])
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic, formatted_str_sections=formatted_str_sections
            )
            conclusion = self.llm.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content="write report conclusion"),
                ]
            )
            self.logger.info(
                "successfully generated the conclusion", length=len(conclusion.content)
            )
            return {"conclusion": conclusion.content}

        except Exception as e:
            self.logger.error("failed to generate the conclusion", error=str(e))
            raise ResearchAnalystException("hmmm...not there exactly", e)

    def finalize_report(self, state: ResearchGraphState):
        """
        finalizes the report and writes the final piece
        """
        content = state.get("content", "")

        try:
            self.logger.info("finalises the report")

            if content.startswith("## Insights"):
                content = content[len("## Insights") :].lstrip()

            sources = None
            if "\n## Sources\n" in content:
                content, sources = content.split("\n## Sources\n", 1)

            final_report = (
                state.get("introduction", "")
                + "\n\n---\n\n"
                + content.strip()
                + "\n\n---\n\n"
                + state.get("conclusion", "")
            )

            if sources:
                final_report += "\n\n## Sources\n" + sources.strip()

            self.logger.info("Voila!")
            return {"final_report": final_report}

        except Exception as e:
            self.logger.error("failed process", error=str(e))
            raise ResearchAnalystException("heartbreak for you bruh!", e)

    def save_report(self, final_report: str, topic: str, format: str = "docx"):
        """
        saves the DOCX or PDF in respective subfolder
        """
        try:
            self.logger.info("saving report", topic=topic, format=format)
            timestamp = datetime.now().strftime("%d%m%Y_%H%M%S")
            safe_topic = re.sub(r'[\\/*?:"<>]', "_", topic)
            base_name = f"{safe_topic.replace(' ', '_')}_{timestamp}"

            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
            root_dir = os.path.join(project_root, "generated_report")

            report_folder = os.path.join(root_dir, base_name)
            os.makedirs(report_folder, exist_ok=True)

            file_path = os.path.join(report_folder, f"{base_name}.{format}")

            if format == "docx":
                self._save_as_docx(final_report, file_path)
            elif format == "pdf":
                self._save_as_pdf(final_report, file_path)
            else:
                raise ValueError("Invalid format. Use 'docx' or 'pdf'.")

            self.logger.info("report saved successfully", path=file_path)
            return file_path

        except Exception as e:
            self.logger.error("bhool e bhora bosonto", error=str(e))
            raise ResearchAnalystException("Akasher thikanay chithi dilam", e)

    def _save_as_docx(self, text: str, file_path: str):
        """
        helps to save the document in docx format
        """
        try:
            doc = Document()
            for line in text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            self.logger.error("cudnt save", path=file_path, error=str(e))
            raise ResearchAnalystException("process failed..please retry", e)

    def _save_as_pdf(self, text: str, file_path: str):
        """
        this function helps creating the pdf
        """
        from textwrap import wrap

        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            left_margin = 80
            right_margin = 80
            usable_width = width - left_margin - right_margin
            top_margin = 70
            bottom_margin = 80
            y = height - top_margin

            normal_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            line_height = 15

            lines = text.split("\n")
            for raw_lines in lines:
                line = raw_lines.strip()
                if not line:
                    y -= line_height
                    continue

                if line.startswith("# "):
                    font = bold_font
                    size = 16
                    line = line[2:].strip()
                elif line.startswith("## "):
                    font = bold_font
                    size = 13
                    line = line[3:].strip()
                else:
                    font = normal_font
                    size = 11

                c.setFont(font, size)
                wrapped_lines = wrap(line, width=int(usable_width / (size * 0.55)))

                for wline in wrapped_lines:
                    if y < bottom_margin:
                        c.showPage()
                        c.setFont(font, size)
                        y = height - top_margin

                    text_width = c.stringWidth(wline, font, size)
                    x = (width - text_width) / 2
                    c.drawString(x, y, wline)
                    y -= line_height

            c.save()
            self.logger.info("centered PDF", path=file_path)

        except Exception as e:
            self.logger.error("PDF save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Please retry", e)

    def build_graph(self):
        """
        constructing the complete graph
        """
        try:
            self.logger.info("building the complete graph")
            builder = StateGraph(ResearchGraphState)
            interview_graph = InterviewGraphBuilder(self.llm, self.tavily_search).build()

            def initiate_all_interviews(state: ResearchGraphState):
                topic = state.get("topic", "Untitled Topic")
                analysts = state.get("analysts", [])
                if not analysts:
                    self.logger.warning("no analysts found, skipping interviews")
                    return END
                return [
                    Send(
                        "conduct_interview",
                        {
                            "analyst": analyst,
                            "messages": [
                                HumanMessage(content=f"So let's discuss about the topic {topic}")
                            ],
                            "max_num_turns": 2,
                            "context": [],
                            "interview": "",
                            "sections": [],
                        },
                    )
                    for analyst in analysts
                ]

            builder.add_node("create_analyst", self.create_analyst)
            builder.add_node("human_feedback", self.human_feedback)
            builder.add_node("conduct_interview", interview_graph)
            builder.add_node("write_report", self.write_report)
            builder.add_node("write_introduction", self.write_introduction)
            builder.add_node("write_conclusion", self.write_conclusion)
            builder.add_node("finalize_report", self.finalize_report)

            builder.add_edge(START, "create_analyst")
            builder.add_edge("create_analyst", "human_feedback")
            builder.add_conditional_edges(
                "human_feedback",
                initiate_all_interviews,
                ["conduct_interview", END],
            )

            builder.add_edge("conduct_interview", "write_report")
            builder.add_edge("conduct_interview", "write_introduction")
            builder.add_edge("conduct_interview", "write_conclusion")
            builder.add_edge(
                ["write_report", "write_introduction", "write_conclusion"],
                "finalize_report",
            )
            builder.add_edge("finalize_report", END)

            graph = builder.compile(
                interrupt_before=["human_feedback"], checkpointer=self.memory
            )
            self.logger.info("Report generation completed successfully")
            return graph

        except Exception as e:
            self.logger.error("error in building graph", error=str(e))
            raise ResearchAnalystException("failed process", e)


if __name__ == "__main__":
    try:
        llm = ModelLoader().load_llm()
        reporter = AutonomousReportGenerator(llm)
        graph = reporter.build_graph()

        topic = "Impact of LLMs on Marketing"
        thread = {"configurable": {"thread_id": "1"}}
        reporter.logger.info("Starting report generation pipeline", topic=topic)

        for _ in graph.stream({"topic": topic, "max_analysts": 3}, thread, stream_mode="values"):
            pass

        state = graph.get_state(thread)
        feedback = input("\n Enter your feedback or press Enter to continue: ").strip()
        graph.update_state(thread, {"human_analyst_feedback": feedback}, as_node="human_feedback")

        for _ in graph.stream(None, thread, stream_mode="values"):
            pass

        final_state = graph.get_state(thread)
        final_report = final_state.values.get("final_report")

        if final_report:
            reporter.logger.info("report generated successfully")
            reporter.save_report(final_report, topic, "docx")
            reporter.save_report(final_report, topic, "pdf")
        else:
            reporter.logger.error("No report generated")

    except Exception as e:
        print(e)
        raise ResearchAnalystException("please check and retry", e)